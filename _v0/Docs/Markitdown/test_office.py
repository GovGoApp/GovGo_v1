"""
Teste de conversão de arquivos do Microsoft Office com MarkItDown
Demonstra conversão de Word, Excel e PowerPoint
"""

from markitdown import MarkItDown
import os
from datetime import datetime

def create_test_excel():
    """Cria um arquivo Excel de teste"""
    try:
        import pandas as pd
        
        # Criar dados de teste
        data = {
            'Nome': ['João Silva', 'Maria Santos', 'Pedro Oliveira', 'Ana Costa'],
            'Idade': [30, 25, 35, 28],
            'Departamento': ['TI', 'RH', 'Vendas', 'Marketing'],
            'Salário': [5000, 4500, 6000, 4800],
            'Data Admissão': ['2020-01-15', '2021-03-20', '2019-11-10', '2022-05-05']
        }
        
        df = pd.DataFrame(data)
        df.to_excel('test_data.xlsx', index=False, sheet_name='Funcionários')
        
        print("✅ Arquivo Excel de teste criado: test_data.xlsx")
        return True
        
    except ImportError:
        print("⚠️  Pandas não instalado. Para criar Excel, instale: pip install pandas openpyxl")
        return False

def test_excel_conversion():
    """Teste de conversão de Excel"""
    md = MarkItDown()
    
    print("=== Teste Conversão Excel ===\n")
    
    if create_test_excel():
        try:
            # Converter Excel
            result = md.convert("test_data.xlsx")
            print("Conversão Excel -> Markdown:")
            print("-" * 40)
            print(result.text_content)
            print("-" * 40)
            
            # Salvar resultado
            with open("output_excel.md", "w", encoding="utf-8") as f:
                f.write(result.text_content)
            
            print("✅ Conversão Excel salva em: output_excel.md\n")
            
        except Exception as e:
            print(f"❌ Erro na conversão Excel: {e}")
            print("Pode precisar instalar: pip install 'markitdown[xlsx]'\n")
        
        finally:
            # Limpar arquivo temporário
            if os.path.exists("test_data.xlsx"):
                os.remove("test_data.xlsx")

def create_test_word():
    """Cria um documento Word de teste"""
    try:
        from docx import Document
        
        doc = Document()
        doc.add_heading('Relatório de Teste', 0)
        
        doc.add_heading('Introdução', level=1)
        doc.add_paragraph('Este é um documento de teste para demonstrar a conversão do MarkItDown.')
        
        doc.add_heading('Dados', level=1)
        doc.add_paragraph('Alguns dados importantes:')
        
        # Adicionar lista
        doc.add_paragraph('Item 1', style='List Bullet')
        doc.add_paragraph('Item 2', style='List Bullet')
        doc.add_paragraph('Item 3', style='List Bullet')
        
        doc.add_heading('Conclusão', level=1)
        doc.add_paragraph('Este documento foi convertido com sucesso.')
        
        doc.save('test_document.docx')
        print("✅ Documento Word de teste criado: test_document.docx")
        return True
        
    except ImportError:
        print("⚠️  python-docx não instalado. Para criar Word, instale: pip install python-docx")
        return False

def test_word_conversion():
    """Teste de conversão de Word"""
    md = MarkItDown()
    
    print("=== Teste Conversão Word ===\n")
    
    if create_test_word():
        try:
            # Converter Word
            result = md.convert("test_document.docx")
            print("Conversão Word -> Markdown:")
            print("-" * 40)
            print(result.text_content)
            print("-" * 40)
            
            # Salvar resultado
            with open("output_word.md", "w", encoding="utf-8") as f:
                f.write(result.text_content)
            
            print("✅ Conversão Word salva em: output_word.md\n")
            
        except Exception as e:
            print(f"❌ Erro na conversão Word: {e}")
            print("Pode precisar instalar: pip install 'markitdown[docx]'\n")
        
        finally:
            # Limpar arquivo temporário
            if os.path.exists("test_document.docx"):
                os.remove("test_document.docx")

def test_url_conversion():
    """Teste de conversão de URL"""
    md = MarkItDown()
    
    print("=== Teste Conversão URL ===\n")
    
    # URLs de teste
    test_urls = [
        "https://httpbin.org/html",  # HTML simples
        "https://jsonplaceholder.typicode.com/posts/1"  # JSON
    ]
    
    for i, url in enumerate(test_urls, 1):
        try:
            print(f"Convertendo URL {i}: {url}")
            result = md.convert(url)
            
            print(f"Conversão URL {i} -> Markdown:")
            print("-" * 40)
            print(result.text_content[:500] + "..." if len(result.text_content) > 500 else result.text_content)
            print("-" * 40)
            
            # Salvar resultado
            with open(f"output_url_{i}.md", "w", encoding="utf-8") as f:
                f.write(result.text_content)
            
            print(f"✅ Conversão URL {i} salva em: output_url_{i}.md\n")
            
        except Exception as e:
            print(f"❌ Erro na conversão URL {i}: {e}\n")

if __name__ == "__main__":
    test_excel_conversion()
    test_word_conversion()
    test_url_conversion()
    print("=== Testes de Office Concluídos ===")
